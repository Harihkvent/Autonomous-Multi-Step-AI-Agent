import os
import re
import docx
import tempfile
from datetime import datetime
from PyPDF2 import PdfReader
from tools.registry import registry
from models import ToolResult

def doc_parser(filepath: str) -> str:
    """Extract text from .docx, .pdf, or .txt files."""
    if not os.path.exists(filepath):
        return f"Error: File not found at {filepath}"
    
    ext = filepath.split('.')[-1].lower()
    text = ""
    try:
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        elif ext == 'pdf':
            reader = PdfReader(filepath)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif ext == 'docx':
            doc = docx.Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs])
            
        if len(text) > 3000:
            text = text[:3000] + "\n...[TRUNCATED]"
        return text if text else "(No text content extracted)"
    except Exception as e:
        return f"Error parsing {filepath}: {str(e)}"

def doc_generator(topic_or_content: str) -> str:
    """Generate a professional .docx document. If input is short, it auto-expands using AI."""
    from core.graph import generate_krutrim_response
    from langchain_core.messages import HumanMessage, SystemMessage
    
    content = topic_or_content
    # If content is short (just a topic), use LLM to generate full content
    if len(content) < 200:
        content = generate_krutrim_response([
            SystemMessage(content="You are a professional document writer. Write a detailed, well-structured document on the given topic. Include sections with headings, bullet points, and comprehensive content. Write at least 500 words."),
            HumanMessage(content=f"Write a detailed document about: {content}")
        ])
    
    doc = docx.Document()
    # Extract a short title
    topic_title = topic_or_content[:80].split('.')[0].strip()
    doc.add_heading(topic_title if len(topic_title) > 3 else "AI Generated Document", 0)
    
    # Smart paragraph insertion — detect markdown headings
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('#'):
            heading_text = line.lstrip('#').strip()
            level = 1 if not line.startswith('##') else 2
            doc.add_heading(heading_text, level=level)
        else:
            doc.add_paragraph(line)
    
    docs_dir = os.path.join(tempfile.gettempdir(), "agent_generated_docs")
    os.makedirs(docs_dir, exist_ok=True)
    filename = f"Generated_Report_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    file_path = os.path.join(docs_dir, filename)
    doc.save(file_path)
    return f"Successfully generated document: {filename} [DOWNLOAD:{filename}]"

def text_writer(prompt: str) -> str:
    """Generate raw text or creative content using the AI model."""
    from core.graph import generate_krutrim_response
    from langchain_core.messages import HumanMessage
    return generate_krutrim_response([HumanMessage(content=prompt)])

# Register tools
registry.register("doc_parser", "Extract text from documents (.pdf, .docx, .txt). Args: filepath", doc_parser)
registry.register("doc_generator", "Generate a .docx report. Args: topic_or_content", doc_generator)
registry.register("text_writer", "Generate structured text, emails, or essays. Args: prompt", text_writer)

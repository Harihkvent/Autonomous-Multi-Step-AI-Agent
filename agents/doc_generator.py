import os
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
import docx
import tempfile
from datetime import datetime

def doc_generator_node(state):
    messages = state.get("messages", [])
    last_message = messages[-1].content
    
    try:
        # Use LLM to generate actual document content
        from core.graph import generate_krutrim_response
        
        topic = last_message.replace("DocGenerator Agent:", "").strip()
        
        generated_content = generate_krutrim_response([
            SystemMessage(content="You are a professional document writer. Write a detailed, well-structured document on the given topic. Include sections with headings, bullet points, and comprehensive content. Write at least 500 words."),
            HumanMessage(content=f"Write a detailed document about: {topic}")
        ])
        
        doc = docx.Document()
        # Extract a title from the topic
        title = topic.split('.')[0].split(',')[0][:80].strip()
        doc.add_heading(title if len(title) > 3 else "AI Generated Document", 0)
        
        # Split content by lines and add paragraphs
        for line in generated_content.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Detect headings (lines that look like section titles)
            if line.startswith('#'):
                heading_text = line.lstrip('#').strip()
                doc.add_heading(heading_text, level=1)
            elif line.startswith('##'):
                heading_text = line.lstrip('#').strip()
                doc.add_heading(heading_text, level=2)
            else:
                doc.add_paragraph(line)
        
        docs_dir = os.path.join(tempfile.gettempdir(), "agent_generated_docs")
        os.makedirs(docs_dir, exist_ok=True)
        filename = f"Generated_Report_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        file_path = os.path.join(docs_dir, filename)
        doc.save(file_path)
        
        msg = AIMessage(content=f"DocGenerator Agent: Successfully generated document: {filename} [DOWNLOAD:{filename}]", name="doc_generator")
    except Exception as e:
        msg = AIMessage(content=f"DocGenerator Agent: Failed to generate document. Error: {e}", name="doc_generator")
        
    return {"messages": [msg], "next": "supervisor"}
